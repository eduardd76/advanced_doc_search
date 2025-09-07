"""
Advanced document processing system with multi-format support
"""

import os
import io
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass
import hashlib
from collections import defaultdict
import chardet
from langdetect import detect

logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Represents a processed document"""
    doc_id: str
    file_path: str
    content: str
    metadata: Dict[str, Any]
    extraction_method: str
    language: Optional[str] = None
    page_count: Optional[int] = None
    word_count: int = 0
    tables: List[Dict[str, Any]] = None
    images: List[Dict[str, Any]] = None
    
    def __post_init__(self):
        if self.tables is None:
            self.tables = []
        if self.images is None:
            self.images = []
        if self.word_count == 0:
            self.word_count = len(self.content.split())

@dataclass
class ExtractionResult:
    """Result of document extraction"""
    success: bool
    content: str
    metadata: Dict[str, Any]
    extraction_method: str
    error: Optional[str] = None

class PDFProcessor:
    """Advanced PDF processing with multiple extraction methods"""
    
    def __init__(self, enable_ocr: bool = True, ocr_languages: List[str] = None):
        self.enable_ocr = enable_ocr
        self.ocr_languages = ocr_languages or ['eng']
        self.extraction_methods = ['pymupdf', 'pdfplumber', 'pdfminer', 'ocr']
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from PDF using multiple fallback methods"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return ExtractionResult(
                success=False,
                content="",
                metadata={},
                extraction_method="none",
                error=f"File not found: {file_path}"
            )
        
        # Try each extraction method in order
        for method in self.extraction_methods:
            try:
                if method == 'pymupdf':
                    result = self._extract_with_pymupdf(file_path)
                elif method == 'pdfplumber':
                    result = self._extract_with_pdfplumber(file_path)
                elif method == 'pdfminer':
                    result = self._extract_with_pdfminer(file_path)
                elif method == 'ocr' and self.enable_ocr:
                    result = self._extract_with_ocr(file_path)
                else:
                    continue
                
                if result.success and len(result.content.strip()) > 50:
                    logger.debug(f"Successfully extracted PDF with {method}: {len(result.content)} chars")
                    return result
                    
            except Exception as e:
                logger.warning(f"PDF extraction method {method} failed for {file_path}: {e}")
                continue
        
        return ExtractionResult(
            success=False,
            content="",
            metadata={},
            extraction_method="failed",
            error="All PDF extraction methods failed"
        )
    
    def _extract_with_pymupdf(self, file_path: Path) -> ExtractionResult:
        """Extract using PyMuPDF (fitz)"""
        try:
            import fitz
        except ImportError:
            raise ImportError("PyMuPDF not available")
        
        doc = fitz.open(str(file_path))
        text_parts = []
        tables = []
        images = []
        
        for page_num, page in enumerate(doc):
            # Extract text
            page_text = page.get_text()
            text_parts.append(page_text)
            
            # Extract tables
            try:
                page_tables = page.find_tables()
                for table in page_tables:
                    table_data = table.extract()
                    tables.append({
                        'page': page_num + 1,
                        'data': table_data,
                        'bbox': table.bbox
                    })
            except:
                pass  # Table extraction is optional
            
            # Extract images info
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                images.append({
                    'page': page_num + 1,
                    'index': img_index,
                    'width': img[2],
                    'height': img[3]
                })
        
        doc.close()
        
        content = "\n".join(text_parts)
        metadata = {
            'page_count': len(text_parts),
            'tables_count': len(tables),
            'images_count': len(images),
            'tables': tables,
            'images': images
        }
        
        return ExtractionResult(
            success=len(content.strip()) > 0,
            content=content,
            metadata=metadata,
            extraction_method="pymupdf"
        )
    
    def _extract_with_pdfplumber(self, file_path: Path) -> ExtractionResult:
        """Extract using pdfplumber"""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber not available")
        
        text_parts = []
        tables = []
        
        with pdfplumber.open(str(file_path)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                
                # Extract tables
                page_tables = page.extract_tables()
                for table in page_tables:
                    tables.append({
                        'page': page_num + 1,
                        'data': table
                    })
        
        content = "\n".join(text_parts)
        metadata = {
            'page_count': len(text_parts),
            'tables_count': len(tables),
            'tables': tables
        }
        
        return ExtractionResult(
            success=len(content.strip()) > 0,
            content=content,
            metadata=metadata,
            extraction_method="pdfplumber"
        )
    
    def _extract_with_pdfminer(self, file_path: Path) -> ExtractionResult:
        """Extract using pdfminer.six"""
        try:
            from pdfminer.high_level import extract_text
            from pdfminer.pdfpage import PDFPage
        except ImportError:
            raise ImportError("pdfminer.six not available")
        
        # Extract text
        text = extract_text(str(file_path))
        
        # Count pages
        page_count = 0
        with open(file_path, 'rb') as fp:
            for page in PDFPage.get_pages(fp):
                page_count += 1
        
        metadata = {'page_count': page_count}
        
        return ExtractionResult(
            success=len(text.strip()) > 0,
            content=text,
            metadata=metadata,
            extraction_method="pdfminer"
        )
    
    def _extract_with_ocr(self, file_path: Path) -> ExtractionResult:
        """Extract using OCR (Tesseract)"""
        try:
            import fitz
            import pytesseract
            from PIL import Image
        except ImportError:
            raise ImportError("OCR dependencies not available")
        
        doc = fitz.open(str(file_path))
        text_parts = []
        
        for page_num, page in enumerate(doc):
            # Convert page to image
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")
            
            # OCR the image
            image = Image.open(io.BytesIO(img_data))
            page_text = pytesseract.image_to_string(
                image, 
                lang='+'.join(self.ocr_languages)
            )
            text_parts.append(page_text)
        
        doc.close()
        
        content = "\n".join(text_parts)
        metadata = {
            'page_count': len(text_parts),
            'ocr_languages': self.ocr_languages
        }
        
        return ExtractionResult(
            success=len(content.strip()) > 0,
            content=content,
            metadata=metadata,
            extraction_method="ocr"
        )

class DOCXProcessor:
    """Process DOCX documents"""
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from DOCX"""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx not available")
        
        try:
            doc = docx.Document(file_path)
            text_parts = []
            tables = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    'table_number': table_num + 1,
                    'data': table_data
                })
            
            content = "\n".join(text_parts)
            metadata = {
                'paragraph_count': len(text_parts),
                'tables_count': len(tables),
                'tables': tables
            }
            
            return ExtractionResult(
                success=len(content.strip()) > 0,
                content=content,
                metadata=metadata,
                extraction_method="python-docx"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata={},
                extraction_method="python-docx",
                error=str(e)
            )

class EPUBProcessor:
    """Process EPUB documents"""
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from EPUB"""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("ebooklib and beautifulsoup4 not available for EPUB processing")
        
        try:
            book = epub.read_epub(file_path)
            text_parts = []
            
            # Extract text from all items
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_body_content(), 'html.parser')
                    text = soup.get_text()
                    if text.strip():
                        text_parts.append(text)
            
            content = "\n".join(text_parts)
            metadata = {
                'item_count': len(text_parts),
                'title': book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else None,
                'author': book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else None,
                'language': book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else None
            }
            
            return ExtractionResult(
                success=len(content.strip()) > 0,
                content=content,
                metadata=metadata,
                extraction_method="ebooklib"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata={},
                extraction_method="ebooklib",
                error=str(e)
            )

class RTFProcessor:
    """Process RTF documents"""
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from RTF"""
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            raise ImportError("striprtf not available for RTF processing")
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            # Extract plain text
            text = rtf_to_text(rtf_content)
            
            # Basic cleaning
            text = self._clean_text(text)
            
            metadata = {
                'file_size': len(rtf_content),
                'extracted_length': len(text)
            }
            
            return ExtractionResult(
                success=len(text.strip()) > 0,
                content=text,
                metadata=metadata,
                extraction_method="striprtf"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata={},
                extraction_method="striprtf", 
                error=str(e)
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        import re
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        return text.strip()

class HTMLProcessor:
    """Process HTML documents"""
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from HTML"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 not available for HTML processing")
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean text
            text = self._clean_text(text)
            
            # Extract metadata
            title = soup.find('title')
            title_text = title.string if title else None
            
            meta_description = soup.find('meta', attrs={'name': 'description'})
            description = meta_description.get('content') if meta_description else None
            
            metadata = {
                'title': title_text,
                'description': description,
                'file_size': len(html_content)
            }
            
            return ExtractionResult(
                success=len(text.strip()) > 0,
                content=text,
                metadata=metadata,
                extraction_method="beautifulsoup4"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata={},
                extraction_method="beautifulsoup4",
                error=str(e)
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        import re
        # Normalize line endings and spaces
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        return text.strip()

class TextProcessor:
    """Process plain text and markdown files"""
    
    def extract_text(self, file_path: str) -> ExtractionResult:
        """Extract text from plain text files"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            # Try to detect encoding
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result.get('encoding', 'utf-8')
            confidence = encoding_result.get('confidence', 0.0)
            
            # Fallback encodings to try
            encodings_to_try = [encoding, 'utf-8', 'latin-1', 'cp1252']
            
            for enc in encodings_to_try:
                try:
                    content = raw_data.decode(enc)
                    
                    # Basic cleaning
                    content = self._clean_text(content)
                    
                    metadata = {
                        'encoding': enc,
                        'encoding_confidence': confidence,
                        'file_size': len(raw_data),
                        'line_count': content.count('\n') + 1
                    }
                    
                    return ExtractionResult(
                        success=len(content.strip()) > 0,
                        content=content,
                        metadata=metadata,
                        extraction_method="text"
                    )
                    
                except UnicodeDecodeError:
                    continue
            
            return ExtractionResult(
                success=False,
                content="",
                metadata={},
                extraction_method="text",
                error="Could not decode text file with any encoding"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                content="",
                metadata={},
                extraction_method="text",
                error=str(e)
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove null characters
        text = text.replace('\x00', '')
        
        return text.strip()

class BoilerplateRemover:
    """Remove boilerplate content from documents"""
    
    def __init__(self):
        # Common boilerplate patterns
        self.boilerplate_patterns = [
            # Headers and footers
            r'Page \d+ of \d+',
            r'©.*?\d{4}.*?All rights reserved',
            r'Confidential and Proprietary',
            r'This document contains proprietary information',
            
            # Navigation elements
            r'Table of Contents',
            r'Index',
            r'Bibliography',
            r'References',
            
            # Repeated elements
            r'Click here to.*?',
            r'For more information visit.*?',
            r'Downloaded from.*?',
        ]
        
        self.compiled_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in self.boilerplate_patterns]
    
    def remove_boilerplate(self, text: str) -> str:
        """Remove boilerplate content"""
        # Remove common boilerplate patterns
        for pattern in self.compiled_patterns:
            text = pattern.sub('', text)
        
        # Remove repeated lines (headers/footers)
        text = self._remove_repeated_lines(text)
        
        # Remove very short paragraphs that are likely boilerplate
        text = self._remove_short_paragraphs(text)
        
        return text
    
    def _remove_repeated_lines(self, text: str) -> str:
        """Remove lines that appear too frequently (likely headers/footers)"""
        lines = text.split('\n')
        line_counts = defaultdict(int)
        
        # Count occurrences of each line
        for line in lines:
            cleaned_line = line.strip()
            if len(cleaned_line) > 5 and len(cleaned_line) < 100:
                line_counts[cleaned_line] += 1
        
        # Remove lines that appear too frequently
        threshold = max(3, len(lines) // 10)  # Appear in >10% of text
        repeated_lines = {line for line, count in line_counts.items() if count > threshold}
        
        # Filter out repeated lines
        filtered_lines = []
        for line in lines:
            if line.strip() not in repeated_lines:
                filtered_lines.append(line)
        
        return '\n'.join(filtered_lines)
    
    def _remove_short_paragraphs(self, text: str) -> str:
        """Remove very short paragraphs that are likely navigation elements"""
        paragraphs = text.split('\n\n')
        filtered_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            # Keep paragraphs that are long enough or contain important content
            if (len(para.split()) >= 5 or 
                any(keyword in para.lower() for keyword in ['figure', 'table', 'chapter', 'section'])):
                filtered_paragraphs.append(para)
        
        return '\n\n'.join(filtered_paragraphs)

class AdvancedDocumentProcessor:
    """
    Advanced document processor that handles multiple formats with:
    - Multi-format extraction
    - Layout-aware processing  
    - Boilerplate removal
    - Language detection
    - Metadata extraction
    """
    
    def __init__(
        self,
        supported_formats: List[str] = None,
        enable_ocr: bool = True,
        ocr_languages: List[str] = None,
        remove_boilerplate: bool = True,
        detect_language: bool = True
    ):
        self.supported_formats = supported_formats or [
            '.pdf', '.docx', '.txt', '.md', '.epub', '.rtf', '.html'
        ]
        self.enable_ocr = enable_ocr
        self.ocr_languages = ocr_languages or ['eng']
        self.remove_boilerplate = remove_boilerplate
        self.detect_language = detect_language
        
        # Initialize processors
        self.pdf_processor = PDFProcessor(enable_ocr, ocr_languages)
        self.docx_processor = DOCXProcessor()
        self.epub_processor = EPUBProcessor()
        self.rtf_processor = RTFProcessor()
        self.html_processor = HTMLProcessor()
        self.text_processor = TextProcessor()
        
        # Initialize boilerplate remover
        self.boilerplate_remover = BoilerplateRemover() if remove_boilerplate else None
    
    def process_document(self, file_path: str, doc_id: Optional[str] = None) -> ProcessedDocument:
        """Process a document and extract content"""
        file_path = Path(file_path)
        
        if doc_id is None:
            doc_id = self._generate_doc_id(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return self._create_error_document(doc_id, str(file_path), "File not found")
        
        # Check file extension
        extension = file_path.suffix.lower()
        if extension not in self.supported_formats:
            logger.warning(f"Unsupported file format: {extension}")
            return self._create_error_document(doc_id, str(file_path), f"Unsupported format: {extension}")
        
        # Extract content based on file type
        try:
            if extension == '.pdf':
                extraction_result = self.pdf_processor.extract_text(str(file_path))
            elif extension == '.docx':
                extraction_result = self.docx_processor.extract_text(str(file_path))
            elif extension == '.epub':
                extraction_result = self.epub_processor.extract_text(str(file_path))
            elif extension == '.rtf':
                extraction_result = self.rtf_processor.extract_text(str(file_path))
            elif extension == '.html':
                extraction_result = self.html_processor.extract_text(str(file_path))
            elif extension in ['.txt', '.md']:
                extraction_result = self.text_processor.extract_text(str(file_path))
            else:
                extraction_result = ExtractionResult(
                    success=False,
                    content="",
                    metadata={},
                    extraction_method="none",
                    error=f"No processor for {extension}"
                )
            
            if not extraction_result.success:
                logger.error(f"Failed to extract content from {file_path}: {extraction_result.error}")
                return self._create_error_document(doc_id, str(file_path), extraction_result.error)
            
            # Post-process content
            content = extraction_result.content
            
            # Remove boilerplate if enabled
            if self.boilerplate_remover and content:
                content = self.boilerplate_remover.remove_boilerplate(content)
            
            # Detect language if enabled
            language = None
            if self.detect_language and len(content.strip()) > 50:
                try:
                    language = detect(content[:1000])  # Use first 1000 chars for detection
                except:
                    language = None
            
            # Create processed document
            metadata = {
                **extraction_result.metadata,
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': extension,
                'extraction_method': extraction_result.extraction_method
            }
            
            processed_doc = ProcessedDocument(
                doc_id=doc_id,
                file_path=str(file_path),
                content=content,
                metadata=metadata,
                extraction_method=extraction_result.extraction_method,
                language=language,
                page_count=extraction_result.metadata.get('page_count'),
                tables=extraction_result.metadata.get('tables', []),
                images=extraction_result.metadata.get('images', [])
            )
            
            logger.debug(f"Processed document {doc_id}: {processed_doc.word_count} words, language: {language}")
            return processed_doc
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return self._create_error_document(doc_id, str(file_path), str(e))
    
    def process_directory(
        self, 
        directory_path: str, 
        recursive: bool = True,
        file_extensions: Optional[List[str]] = None
    ) -> List[ProcessedDocument]:
        """Process all supported documents in a directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        if file_extensions is None:
            file_extensions = self.supported_formats
        
        # Find all files
        pattern = "**/*" if recursive else "*"
        all_files = directory_path.glob(pattern)
        
        # Filter by extension
        filtered_files = [
            f for f in all_files 
            if f.is_file() and f.suffix.lower() in file_extensions
        ]
        
        logger.info(f"Processing {len(filtered_files)} files from {directory_path}")
        
        # Process files
        processed_docs = []
        for file_path in filtered_files:
            try:
                doc = self.process_document(str(file_path))
                if doc.content.strip():  # Only add documents with content
                    processed_docs.append(doc)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        logger.info(f"Successfully processed {len(processed_docs)} documents")
        return processed_docs
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """Generate unique document ID"""
        # Use file path hash for consistent IDs
        path_str = str(file_path.absolute())
        return hashlib.md5(path_str.encode()).hexdigest()[:12]
    
    def _create_error_document(self, doc_id: str, file_path: str, error: str) -> ProcessedDocument:
        """Create document for processing errors"""
        return ProcessedDocument(
            doc_id=doc_id,
            file_path=file_path,
            content="",
            metadata={'error': error, 'file_path': file_path},
            extraction_method="error",
            word_count=0
        )
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return self.supported_formats.copy()
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get processor statistics"""
        return {
            'supported_formats': self.supported_formats,
            'enable_ocr': self.enable_ocr,
            'ocr_languages': self.ocr_languages,
            'remove_boilerplate': self.remove_boilerplate,
            'detect_language': self.detect_language
        }